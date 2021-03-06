#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import *
from docx.shared import Inches
from docx.oxml.shared import *
from docx.enum.text import WD_BREAK
from docxtpl import DocxTemplate
from Helpers import *


class WordHandler:
    HEADING_1 = 1
    HEADING_2 = 2
    HEADING_3 = 3
    HEADING_4 = 4
    UTF8 = 'utf8'

    word_document = None

    def __init__(self, template_path=None):
        if template_path:
            self.word_document = DocxTemplate(template_path)
        else:
            self.word_document = Document()

    def add_heading(self, heading):
        """NOT CURRENTLY IN USE"""
        self.word_document.add_heading(heading, level=self.HEADING_1)

    def add_topic(self, topic):
        """NOT CURRENTLY IN USE"""
        self.word_document.add_heading(topic['title'], level=self.HEADING_2)

    def add_document(self, document, input_handler):
        """NOT CURRENTLY IN USE"""
        document_id = document["id"]
        # Title
        self.word_document.add_heading(document['title'], level=self.HEADING_3)
        # Description
        self.word_document.add_paragraph(document['description'])
        # Field table
        document_fields = input_handler.get_field_list_by_document_id(document_id)
        table = self.word_document.add_table(0, 2)

        table.autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(4.5)

        # -- Status

        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run('Status:').italic = True
        cells[1].text = input_handler.get_status_name_by_id(document['statusId'])

        # -- HIS number
        if (document['hisNumber']):

            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run('HIS-nummer:').italic = True
            cells[1].text = document['hisNumber']

        # -- Document fields
        for document_field in document_fields:
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(document_field['name'] + ':').italic = True
            cells[1].text = document_field['value']

        # Target groups
        target_groups_table = self.word_document.add_table(0, 3)

        target_groups_table.autofit = False
        target_groups_table.columns[0].width = Inches(1.5)
        target_groups_table.columns[1].width = Inches(3)
        target_groups_table.columns[2].width = Inches(1.5)

        mandatory_dict = input_handler.get_mandatory_dict_on_document_id(document_id)

        for mandataory_id, mandataory in mandatory_dict.iteritems():
            row_number = 1

            for target_group in input_handler.get_target_groups_by_mandatory_id_and_document_id(mandataory_id, document_id):
                cells = target_groups_table.add_row().cells
                if row_number == 1:     # Add mandatory name on first iteration
                    cells[0].paragraphs[0].add_run(mandataory['name'] + ':').italic = True
                else:
                    cells[0].text = ''

                cells[1].text = input_handler.get_target_group_by_id(target_group['targetGroupId'])['name']  # Name of target group
                cells[2].text = input_handler.get_action_name_by_id(target_group['actionId'])

                row_number += 1

            if input_handler.get_target_group_legal_bases_by_document_id(document_id) and mandataory_id == '1':
                text = '• Hjemmel:'
                cells = target_groups_table.add_row().cells
                cells[0].paragraphs[0].add_run(text.decode(self.UTF8)).italic = True
                cells[1].text = input_handler.get_target_group_legal_bases_by_document_id(document_id)
            elif input_handler.get_decided_by_by_document_id(document_id) and mandataory_id == '1':
                text = '• Erstattes av:'
                cells = target_groups_table.add_row().cells
                cells[0].paragraphs[0].add_run(text.decode(self.UTF8)).italic = True
                cells[1].text = input_handler.get_decided_by_by_document_id(document_id)

            notice = input_handler.get_mandatory_notice_by_mandatory_id_and_document_id(mandataory_id, document_id)
            if notice:
                text = '• Merknad:'
                cells = target_groups_table.add_row().cells
                cells[0].paragraphs[0].add_run(text.decode(self.UTF8)).italic = True
                cells[1].text = notice

        # Sections (headings)
        heading_dict = input_handler.get_heading_dict_by_document_id(document_id)
        for heading_id, heading in heading_dict.iteritems():
            self.word_document.add_heading(heading['name'], level=self.HEADING_4)
            self.word_document.add_paragraph(
                input_handler.get_heading_content_by_heading_id_and_document_id(heading_id, document_id)['text']
            )

        # Links
        link_category_dict = input_handler.get_link_category_dict_by_document_id(document_id)

        for link_category_id, link_category in link_category_dict.iteritems():  # Loop through all link categories of the document
            self.word_document.add_heading(link_category['name'], level=self.HEADING_4)
            for link in input_handler.get_links_by_link_category_id_and_document_id(link_category_id, document_id): # for each link in current category
                link_paragraph = self.word_document.add_paragraph(style='ListBullet') #TODO: deprecated, check if therese is a new way to do this.
                self.__add_hyperlink(link_paragraph, link['url'], link['text'])   #TODO: Make sure url starts with http/https else, looking for file on disk.

        # Contact address
        self.word_document.add_heading('Kontaktadresse', level=self.HEADING_4)
        self.word_document.add_paragraph(input_handler.get_contact_address_name_by_document_id(document_id))

    def insert_hyper_links(self):
        """
        Loops through the document an inserts links for [[url|text|website.com]] using python docx (docx).
        :return:
        """
        for paragraph in self.word_document.paragraphs:
            content = self.__get_substring_between(paragraph.text, '[[', ']]')
            if len(content) > 0:
                url_content_array = words = content.split("||")
                type = url_content_array[0]
                if type == TemplateElements.URL:
                    text = url_content_array[1]
                    url = url_content_array[2]
                    paragraph.text = ''
                    self.__add_hyperlink(paragraph, url, text)

    def insert_new_page(self):
        """
        Loops through the document an inserts page break for [[newpage]] using python docx (docx).
        :return:
        """
        for paragraph in self.word_document.paragraphs:
            content = self.__get_substring_between(paragraph.text, '[[', ']]')
            if len(content) > 0:
                url_content_array  = content.split("||")
                type = url_content_array[0]
                if type == TemplateElements.NEW_PAGE:
                    paragraph.text = ''
                    paragraph.add_run().add_break(WD_BREAK.PAGE)

    def __get_substring_between(self, string, first_substring, last_substring):
        """
        Returns a a substring between two specified substring.
        E.g. get_substring_between('first second, third', 'first', 'third') => ' second '.
        :param string: the string you want to return a substring from
        :param first_substring:
        :param last_substring:
        :return:
        """
        try:
            start = string.index(first_substring) + len(first_substring)
            end = string.index(last_substring, start)
            return string[start:end]
        except Exception:
            return ''


    def save_word_document(self, file_path):
       """
        Saves Word document to disk
       :param file_path: e.g. 'c:/document.docx'
       :return:
       """
       self.word_document.save(file_path)

    # Sorce: https://github.com/python-openxml/python-docx/issues/74
    # Credit: https://github.com/rushton3179
    def __add_hyperlink(self, paragraph, url, text):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: A Run object containing the hyperlink
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id, )
        hyperlink.set(qn('w:history'), '1')

        # Create a w:r element
        new_run = OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = OxmlElement('w:rPr')

        # Create a w:rStyle element, note this currently does not add the hyperlink style as its not in
        # the default template, I have left it here in case someone uses one that has the style in it
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')

        # Join all the xml elements together add add the required text to the w:r element
        rPr.append(rStyle)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        # r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return r